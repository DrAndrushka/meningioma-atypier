"""Word tables the journal will accept, with its rules enforced in code.

Four rules decide whether a table is typeset as submitted or sent back, and all
four are checked here rather than trusted to the eye:

**No borders.** A table's borders normally live in its *style definition*, not
in the table itself — a ``Table Grid`` table carries no border element at all
and still prints a full grid. Deleting elements from the table XML therefore
proves nothing. :func:`strip_borders` writes an explicit all-``none`` block as a
direct table property, which overrides the style, and :func:`visible_borders`
counts the edges that would actually print, style included, so a test can assert
zero against what the reader sees.

**Under 6.5 inches.** A table wider than the text block gets set broadside —
rotated ninety degrees on the page — which makes it unreadable next to the text
that discusses it. The width is asserted, not hoped for.

**Body type, double spaced.** No shrinking a table to fit by dropping to 9 pt.
If it does not fit at body size, it has too many columns.

**A `Note:—` block underneath.** Abbreviations defined, alphabetically, in the
journal's own pattern: ``AUC indicates area under the curve; IQR, interquartile
range.`` The note travels with the table because a footnote in a separate file
stops matching the table within one revision.

One table per file, each on its own page. Numbers arrive already formatted —
this module writes strings and never rounds, so a value cannot be shown to two
decimals here and three somewhere else.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable, Sequence

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# The journal's text block. A table wider than this is set broadside.
MAX_WIDTH_IN = 6.5

# Cell widths are stored in twips — twentieths of a point, 1440 to the inch —
# and the conversion rounds, usually *up*. A six-column table laid out to exactly
# 6.5 in came back measuring 6.5014: compliant as floats, over the limit in the
# file. Everything is laid out to this smaller target so the rounding has
# somewhere to go, and the hard limit is checked against the stored value.
LAYOUT_WIDTH_IN = 6.4
TWIPS_PER_INCH = 1440
BODY_PT = 12.0
LINE_SPACING = 2.0
NOTE_PREFIX = "Note:—"          # em dash, no space, per the journal pattern

# Abbreviations this phase may use in a table, defined once so two tables cannot
# gloss the same term differently.
GLOSSARY = {
    "ADC": "apparent diffusion coefficient",
    "AUC": "area under the receiver operating characteristic curve",
    "CI": "confidence interval",
    "IQR": "interquartile range",
    "LR+": "positive likelihood ratio",
    "LR−": "negative likelihood ratio",
    "MICE": "multiple imputation by chained equations",
    "NPV": "negative predictive value",
    "OR": "odds ratio",
    "PPV": "positive predictive value",
    "SD": "standard deviation",
    "VIF": "variance inflation factor",
}


class DocxContractError(Exception):
    """A table breaks a rule the journal enforces."""


_EDGES = ("top", "left", "bottom", "right", "insideH", "insideV")
_INVISIBLE = {"none", "nil"}


def strip_borders(table) -> None:
    """Force every edge off, overriding whatever the table's style says.

    Deleting border elements from the table XML is not enough, and believing it
    is was a real defect here. A table's borders normally live in its *style
    definition*, not in the table itself: a ``Table Grid`` table carries no
    ``tblBorders`` element at all and still prints a full grid. Deleting nothing
    and reporting success is the worst possible outcome for the one rule the
    journal is strictest about.

    So an explicit all-``none`` border block is written as a direct table
    property, which takes precedence over the style.
    """
    for element in list(table._tbl.iter()):
        tag = element.tag
        if tag.endswith("}tblBorders") or tag.endswith("}tcBorders"):
            element.getparent().remove(element)

    borders = OxmlElement("w:tblBorders")
    for edge in _EDGES:
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "none")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        borders.append(element)
    table._tbl.tblPr.append(borders)


def visible_borders(table) -> int:
    """Edges that would actually print, style included.

    Counts what a reader would see, not what the XML happens to contain. An
    explicit border block on the table wins; with none, the style's block is
    what governs, and that is where a stray grid hides.
    """
    def _count(block) -> int:
        return sum(1 for edge in _EDGES
                   for element in block.findall(qn(f"w:{edge}"))
                   if (element.get(qn("w:val")) or "").lower() not in _INVISIBLE)

    own = table._tbl.find(qn("w:tblPr"))
    block = own.find(qn("w:tblBorders")) if own is not None else None
    if block is not None:
        return _count(block)
    try:
        style_block = table.style.element.find(
            qn("w:tblPr")).find(qn("w:tblBorders"))
    except AttributeError:
        return 0
    return _count(style_block) if style_block is not None else 0


def stored_inches(width_in: float) -> float:
    """The width Word will actually record, after rounding to whole twips.

    Checking the float that goes in is not the same as checking the value that
    comes out: the rounding is usually upward, so a layout that sums to exactly
    the limit lands just over it in the file.
    """
    return round(width_in * TWIPS_PER_INCH) / TWIPS_PER_INCH


def _set_cell(cell, text: str, *, bold: bool = False,
              align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = LINE_SPACING
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(BODY_PT)


def abbreviations_used(text: str) -> list[str]:
    """Glossary terms that actually appear, in the order the journal wants.

    Only terms present in the table are defined. A note that glosses "PPV" in a
    table with no predictive values tells the reader the note was not written
    for this table.
    """
    # A trailing letter of either case disqualifies the match, so "OR" does not
    # fire inside "ORDER". A plural "s" is allowed, so "AUCs" still counts.
    found = [term for term in GLOSSARY
             if re.search(rf"(?<![A-Za-z]){re.escape(term)}s?(?![A-Za-z])", text)]
    return sorted(found, key=str.lower)


def abbreviation_note(text: str) -> str:
    """The journal's pattern: first term takes 'indicates', the rest a comma."""
    used = abbreviations_used(text)
    if not used:
        return ""
    first, *rest = used
    parts = [f"{first} indicates {GLOSSARY[first]}"]
    parts += [f"{term}, {GLOSSARY[term]}" for term in rest]
    return "; ".join(parts) + "."


def add_table(document: Document, frame: pd.DataFrame, *, title: str,
              note: str = "", index_header: str = "",
              column_widths: Sequence[float] | None = None) -> None:
    """Write one table, its title above and its `Note:—` block below.

    ``frame`` must already hold strings. Formatting numbers here would put a
    second rounding rule in the codebase, free to disagree with the one that
    produced the figures.
    """
    heading = document.add_paragraph()
    heading.paragraph_format.line_spacing = LINE_SPACING
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(BODY_PT)

    n_cols = len(frame.columns) + 1
    table = document.add_table(rows=len(frame) + 1, cols=n_cols)
    table.autofit = False
    strip_borders(table)

    if column_widths is None:
        first = LAYOUT_WIDTH_IN * 0.34
        rest = (LAYOUT_WIDTH_IN - first) / max(1, n_cols - 1)
        column_widths = [first] + [rest] * (n_cols - 1)
    total = sum(stored_inches(w) for w in column_widths)
    if total > MAX_WIDTH_IN:
        raise DocxContractError(
            f"Table is {total:.3f} in wide; over {MAX_WIDTH_IN} in the journal "
            "sets it broadside. Drop a column or shorten the labels.")
    for i, width in enumerate(column_widths):
        for row in table.rows:
            row.cells[i].width = Inches(width)

    _set_cell(table.rows[0].cells[0], index_header, bold=True)
    for j, name in enumerate(frame.columns, start=1):
        _set_cell(table.rows[0].cells[j], str(name), bold=True)
    for i, (label, row) in enumerate(frame.iterrows(), start=1):
        _set_cell(table.rows[i].cells[0], str(label))
        for j, value in enumerate(row, start=1):
            _set_cell(table.rows[i].cells[j], str(value))

    body = " ".join([index_header, *map(str, frame.columns),
                     *map(str, frame.index.astype(str)),
                     *frame.astype(str).to_numpy().ravel(), note])
    footnote = " ".join(part for part in
                        [NOTE_PREFIX + note.removeprefix(NOTE_PREFIX).lstrip()
                         if note else NOTE_PREFIX.rstrip("—"),
                         abbreviation_note(body)] if part).strip()
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = LINE_SPACING
    note_run = paragraph.add_run(footnote)
    note_run.font.size = Pt(BODY_PT)


def new_document() -> Document:
    """A portrait document at body size, double spaced throughout."""
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    style = document.styles["Normal"]
    style.font.size = Pt(BODY_PT)
    style.paragraph_format.line_spacing = LINE_SPACING
    return document


def write_tables(path: Path | str, tables: Iterable[dict]) -> Path:
    """One file, one table per page. Each entry is the kwargs of `add_table`."""
    document = new_document()
    for i, spec in enumerate(tables):
        if i:
            document.add_page_break()
        add_table(document, **spec)
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    return path


def audit(path: Path | str) -> dict[str, object]:
    """Read a written file back and check it against the journal's rules.

    Reading the file rather than the object in memory: what matters is what a
    production editor opens, and a check that never touches the saved bytes can
    pass while the file on disk is wrong.
    """
    document = Document(str(path))
    widths = []
    for table in document.tables:
        cells = table.rows[0].cells
        widths.append(sum((c.width.inches if c.width else 0.0) for c in cells))
    sizes = {run.font.size.pt for para in document.paragraphs
             for run in para.runs if run.font.size}
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    sizes |= {r.font.size.pt for r in para.runs if r.font.size}
    return {
        "tables": len(document.tables),
        "visible_borders": sum(visible_borders(t) for t in document.tables),
        "max_width_in": max(widths) if widths else 0.0,
        "font_sizes_pt": sorted(sizes),
        "landscape": any(s.orientation == WD_ORIENT.LANDSCAPE
                         for s in document.sections),
        "has_note": any(p.text.startswith(NOTE_PREFIX)
                        for p in document.paragraphs),
    }
