"""Word tables — the journal's rules checked against the saved file, not intent."""
from __future__ import annotations

import pandas as pd
import pytest
from docx import Document
from docx.shared import Inches

from heavy_machinery.config import load as _load_config  # noqa: F401  (sys.path)

import docx_tables as dx


def _frame(cols: int = 3, rows: int = 4) -> pd.DataFrame:
    return pd.DataFrame(
        [[f"v{r}{c}" for c in range(cols)] for r in range(rows)],
        index=[f"Criterion {r}" for r in range(rows)],
        columns=[f"Measure {c}" for c in range(cols)])


def _write(tmp_path, frame=None, **over):
    spec = dict(frame=frame if frame is not None else _frame(),
                title="Table. Demo", note="A note.", index_header="Criterion")
    spec.update(over)
    return dx.write_tables(tmp_path / "table.docx", [spec])


# --- no borders ------------------------------------------------------------
def test_borders_live_in_the_style_so_they_are_overridden_not_deleted():
    """The table XML alone proves nothing — a grid style prints a grid without
    a single ``tblBorders`` element."""
    document = Document()
    table = document.add_table(rows=2, cols=2, style="Table Grid")
    assert "tblBorders" not in table._tbl.xml      # nothing to delete
    assert dx.visible_borders(table) > 0           # yet it would print a grid

    dx.strip_borders(table)
    assert dx.visible_borders(table) == 0
    assert "tblBorders" in table._tbl.xml          # an explicit all-none block


# --- width -----------------------------------------------------------------
def test_widths_are_measured_in_the_units_word_actually_stores():
    """Twips, not EMUs. 0.858 in is recorded as 0.858333 — rounded up, which is
    how 6.5 in of columns once saved as 6.5014."""
    assert dx.stored_inches(0.858) == pytest.approx(1236 / 1440)
    assert dx.stored_inches(0.858) > 0.858

    assert dx.LAYOUT_WIDTH_IN < dx.MAX_WIDTH_IN

    first = dx.MAX_WIDTH_IN * 0.34
    rest = (dx.MAX_WIDTH_IN - first) / 5
    widths = [first] + [rest] * 5              # six columns, as the real table has
    assert sum(widths) == pytest.approx(dx.MAX_WIDTH_IN)
    assert sum(dx.stored_inches(w) for w in widths) > dx.MAX_WIDTH_IN


def test_an_oversized_column_spec_is_refused(tmp_path):
    with pytest.raises(dx.DocxContractError, match="broadside"):
        _write(tmp_path, column_widths=[2.0, 2.0, 2.0, 2.0])


# --- what the saved file actually contains ---------------------------------
def test_the_saved_file_obeys_every_layout_rule(tmp_path):
    """Measured on disk: EMU rounding once pushed a compliant layout over."""
    report = dx.audit(_write(tmp_path))
    assert report["visible_borders"] == 0
    assert report["max_width_in"] <= dx.MAX_WIDTH_IN
    assert report["font_sizes_pt"] == [dx.BODY_PT]
    assert report["landscape"] is False

    # A wider table stays under the limit once stored, too.
    assert dx.audit(_write(tmp_path, frame=_frame(cols=5)))[
        "max_width_in"] <= dx.MAX_WIDTH_IN

    document = Document(str(_write(tmp_path)))
    spacings = {p.paragraph_format.line_spacing for p in document.paragraphs
                if p.text.strip()}
    assert spacings == {dx.LINE_SPACING}


# --- the note --------------------------------------------------------------
def test_the_note_is_written_beneath_the_table(tmp_path):
    document = Document(str(_write(tmp_path)))
    texts = [p.text for p in document.paragraphs]
    note = next(t for t in texts if t.startswith(dx.NOTE_PREFIX))
    assert note.startswith("Note:—A note.")


def test_only_the_abbreviations_that_appear_are_defined():
    """A note glossing PPV in a table without one was not written for it, and
    'OR' must not fire on 'ORDER' — an uppercase letter after it disqualifies."""
    used = dx.abbreviations_used("AUC 0.63 with a 95% CI and an IQR")
    assert used == ["AUC", "CI", "IQR"]
    assert "PPV" not in used

    assert "OR" not in dx.abbreviations_used("Ordered by ORDER of magnitude")
    assert "OR" in dx.abbreviations_used("the OR was 3.6")
    assert "AUC" in dx.abbreviations_used("the two AUCs were compared")

    note = dx.abbreviation_note("SD and ADC and AUC")
    assert note.startswith("ADC indicates apparent diffusion coefficient; ")
    assert note.endswith("SD, standard deviation.")

    assert dx.abbreviation_note("plain words only") == ""
    assert len(set(dx.GLOSSARY.values())) == len(dx.GLOSSARY)


# --- structure -------------------------------------------------------------
def test_the_table_keeps_its_header_index_and_values_verbatim(tmp_path):
    """A second rounding rule here would be free to disagree with the figures."""
    document = Document(str(_write(tmp_path)))
    header = [c.text for c in document.tables[0].rows[0].cells]
    assert header == ["Criterion", "Measure 0", "Measure 1", "Measure 2"]

    first = [row.cells[0].text for row in document.tables[0].rows[1:]]
    assert first == [f"Criterion {r}" for r in range(4)]

    frame = pd.DataFrame([["0.7241"]], index=["x"], columns=["y"])
    document = Document(str(_write(tmp_path, frame=frame)))
    assert document.tables[0].rows[1].cells[1].text == "0.7241"


def test_each_table_starts_on_its_own_page(tmp_path):
    path = dx.write_tables(tmp_path / "two.docx", [
        dict(frame=_frame(), title="Table 1", note="", index_header="A"),
        dict(frame=_frame(), title="Table 2", note="", index_header="A")])
    document = Document(str(path))
    breaks = sum(1 for p in document.paragraphs
                 for r in p.runs if "w:br" in r._element.xml)
    assert breaks >= 1
    assert len(document.tables) == 2


# --- against the real scorecard -------------------------------------------
def test_the_real_criteria_table_passes_every_rule(tmp_path, real_cohort):
    import bend_location as bl
    import collinearity as co
    import criteria as cr
    import dichotomy as di
    import eligibility as el
    import imputation as imp
    import models as mo
    import nonlinearity as nl
    import scorecard as sc
    import separation as sep
    import wobble as wb

    frozen = {"adc_value": 0.72, "max_diameter_cm": 3.81, "tumor_volume": 15.1,
              "edema_volume_cm3": 4.76, "edema_index": 0.0617}
    fits = nl.fit_all(real_cohort)
    separation = sep.separation_table(real_cohort)
    bend = bl.bend_table(real_cohort, fits=fits)
    eligible = el.eligible(el.carry_forward(separation, bend))
    wobble, _ = wb.wobble_table(real_cohort, eligible, n_boot=60, frozen=frozen)
    long = sc.scorecard_long(
        eligible, separation=separation, bend=bend,
        agreement=cr.agreement(cr.criteria_table(real_cohort, eligible),
                               real_cohort),
        wobble=wobble,
        imputation=imp.per_draw_cutpoints(imp.load_draws("output"), eligible,
                                          frozen=frozen),
        dichotomy=di.dichotomy_table(real_cohort, eligible, frozen),
        pairs=co.correlated_pairs(*co.spearman_matrix(real_cohort)),
        coefficients=mo.compare_sets(real_cohort, cutpoints=frozen,
                                     n_boot=40)[1])

    path = dx.write_tables(tmp_path / "criteria.docx", [dict(
        frame=sc.scorecard_wide(long),
        title="Table. Criterion-based appraisal of candidate cut-points",
        note=sc.footnote(), index_header="Criterion")])
    report = dx.audit(path)
    assert report["visible_borders"] == 0
    assert report["max_width_in"] <= dx.MAX_WIDTH_IN
    assert report["font_sizes_pt"] == [dx.BODY_PT]
    assert report["landscape"] is False
    assert report["has_note"]
